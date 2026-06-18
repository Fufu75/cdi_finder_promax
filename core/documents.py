"""Rendu .docx ATS-compatible : une colonne, titres standards, pas de tableau."""
import re
from datetime import datetime
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ── Palette sobre ────────────────────────────────────────────────────────────
NOIR = RGBColor(0x1A, 0x1A, 0x1A)
GRIS = RGBColor(0x55, 0x55, 0x55)
BLEU = RGBColor(0x1F, 0x4E, 0x79)


# ── Helpers de mise en forme ──────────────────────────────────────────────────

def _set_font(run, size: int, bold=False, color=None, italic=False):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color or NOIR


def _add_horizontal_rule(doc: Document):
    """Ligne de séparation fine sous un titre de section."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "1F4E79")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _section_titre(doc: Document, texte: str):
    _add_horizontal_rule(doc)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(texte.upper())
    _set_font(run, 10, bold=True, color=BLEU)


def _set_marges(doc: Document):
    for section in doc.sections:
        section.top_margin = Cm(1.8)
        section.bottom_margin = Cm(1.8)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)


# ── En-tête ───────────────────────────────────────────────────────────────────

def _en_tete(doc: Document, identite: dict, titre_cv: str):
    # Nom
    p_nom = doc.add_paragraph()
    p_nom.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_nom.paragraph_format.space_after = Pt(2)
    run = p_nom.add_run(
        f"{identite['prenom'].upper()} {identite['nom'].upper()}"
    )
    _set_font(run, 20, bold=True, color=BLEU)

    # Titre du poste ciblé
    p_titre = doc.add_paragraph()
    p_titre.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_titre.paragraph_format.space_after = Pt(4)
    run = p_titre.add_run(titre_cv)
    _set_font(run, 11, bold=False, color=GRIS, italic=True)

    # Coordonnées sur une ligne
    parties = [
        identite.get("ville", ""),
        identite.get("telephone", ""),
        identite.get("email", ""),
    ]
    if identite.get("linkedin"):
        parties.append(identite["linkedin"])
    if identite.get("github"):
        parties.append(identite["github"])

    p_coord = doc.add_paragraph()
    p_coord.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_coord.paragraph_format.space_after = Pt(6)
    run = p_coord.add_run("  |  ".join(p for p in parties if p))
    _set_font(run, 9, color=GRIS)


# ── Corps du CV ───────────────────────────────────────────────────────────────

def _accroche(doc: Document, texte: str):
    _section_titre(doc, "Profil")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(texte)
    _set_font(run, 10)


def _experiences(doc: Document, experiences: list):
    _section_titre(doc, "Expériences professionnelles")
    for exp in experiences:
        # Ligne poste | entreprise | lieu | période
        p_head = doc.add_paragraph()
        p_head.paragraph_format.space_before = Pt(6)
        p_head.paragraph_format.space_after = Pt(1)
        run_poste = p_head.add_run(exp["intitule"])
        _set_font(run_poste, 10, bold=True)
        run_sep = p_head.add_run(
            f"  –  {exp['entreprise']}  |  {exp['lieu']}  |  {exp['periode']}"
        )
        _set_font(run_sep, 9, color=GRIS)

        for mission in exp.get("missions", []):
            p_m = doc.add_paragraph(style="List Bullet")
            p_m.paragraph_format.left_indent = Cm(0.4)
            p_m.paragraph_format.space_after = Pt(1)
            run = p_m.add_run(mission)
            _set_font(run, 9.5)


def _formations(doc: Document, formations: list):
    _section_titre(doc, "Formation")
    for f in formations:
        p_head = doc.add_paragraph()
        p_head.paragraph_format.space_before = Pt(6)
        p_head.paragraph_format.space_after = Pt(1)
        run_diplome = p_head.add_run(f["diplome"])
        _set_font(run_diplome, 10, bold=True)
        run_sep = p_head.add_run(
            f"  –  {f['etablissement']}  |  {f['lieu']}  |  {f['periode']}"
        )
        _set_font(run_sep, 9, color=GRIS)
        if f.get("description"):
            p_desc = doc.add_paragraph()
            p_desc.paragraph_format.left_indent = Cm(0.4)
            p_desc.paragraph_format.space_after = Pt(1)
            run = p_desc.add_run(f["description"])
            _set_font(run, 9.5, italic=True, color=GRIS)


def _projets(doc: Document, projets: list):
    if not projets:
        return
    _section_titre(doc, "Projets")
    for proj in projets:
        p_head = doc.add_paragraph()
        p_head.paragraph_format.space_before = Pt(6)
        p_head.paragraph_format.space_after = Pt(1)
        run_titre = p_head.add_run(proj["titre"])
        _set_font(run_titre, 10, bold=True)
        run_tech = p_head.add_run(f"  –  {proj['technologies']}")
        _set_font(run_tech, 9, color=GRIS)
        p_desc = doc.add_paragraph()
        p_desc.paragraph_format.left_indent = Cm(0.4)
        p_desc.paragraph_format.space_after = Pt(1)
        run = p_desc.add_run(proj["description"])
        _set_font(run, 9.5)


def _competences(doc: Document, competences: dict):
    _section_titre(doc, "Compétences")
    labels = {
        "techniques": "Langages & outils",
        "cloud": "Cloud",
        "autres": "Autres",
    }
    for cle, label in labels.items():
        valeurs = competences.get(cle, [])
        if not valeurs:
            continue
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(1)
        run_label = p.add_run(f"{label} : ")
        _set_font(run_label, 9.5, bold=True)
        run_val = p.add_run(", ".join(valeurs))
        _set_font(run_val, 9.5)


def _langues(doc: Document, langues: list):
    _section_titre(doc, "Langues")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run("  |  ".join(langues))
    _set_font(run, 9.5)


# ── Lettre de motivation ──────────────────────────────────────────────────────

def _en_tete_lettre(doc: Document, identite: dict, offre: dict):
    # Expéditeur
    p = doc.add_paragraph()
    run = p.add_run(f"{identite['prenom']} {identite['nom']}")
    _set_font(run, 10, bold=True)
    for ligne in [identite.get("ville",""), identite.get("telephone",""), identite.get("email","")]:
        if ligne:
            p2 = doc.add_paragraph()
            run2 = p2.add_run(ligne)
            _set_font(run2, 10)
            p2.paragraph_format.space_after = Pt(0)

    doc.add_paragraph()

    # Date
    p_date = doc.add_paragraph()
    p_date.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_d = p_date.add_run(
        datetime.now().strftime("Paris, le %d %B %Y")
    )
    _set_font(run_d, 10, italic=True, color=GRIS)

    doc.add_paragraph()

    # Objet
    p_objet = doc.add_paragraph()
    run_o = p_objet.add_run(
        f"Objet : Candidature – {offre.get('poste','Poste')} "
        f"chez {offre.get('entreprise','l\'entreprise')}"
    )
    _set_font(run_o, 10, bold=True)

    doc.add_paragraph()


# ── API publique ──────────────────────────────────────────────────────────────

def creer_cv(cv_data: dict, identite: dict, chemin_sortie: str) -> str:
    """Génère le fichier .docx du CV et retourne le chemin."""
    doc = Document()
    _set_marges(doc)

    # Supprime le style de liste par défaut si absent
    _en_tete(doc, identite, cv_data.get("titre_cv", ""))

    if cv_data.get("resume"):
        _accroche(doc, cv_data["resume"])

    if cv_data.get("experiences"):
        _experiences(doc, cv_data["experiences"])

    if cv_data.get("formations"):
        _formations(doc, cv_data["formations"])

    if cv_data.get("projets"):
        _projets(doc, cv_data["projets"])

    if cv_data.get("competences"):
        _competences(doc, cv_data["competences"])

    if cv_data.get("langues"):
        _langues(doc, cv_data["langues"])

    Path(chemin_sortie).parent.mkdir(parents=True, exist_ok=True)
    doc.save(chemin_sortie)
    return chemin_sortie


def creer_lettre(texte_lettre: str, identite: dict, offre_analysee: dict, chemin_sortie: str) -> str:
    """Génère le fichier .docx de la lettre de motivation et retourne le chemin."""
    doc = Document()
    _set_marges(doc)

    _en_tete_lettre(doc, identite, offre_analysee)

    for paragraphe in texte_lettre.split("\n\n"):
        paragraphe = paragraphe.strip()
        if not paragraphe:
            continue
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.first_line_indent = Cm(1)
        run = p.add_run(paragraphe)
        _set_font(run, 11)

    # Formule de politesse + signature
    doc.add_paragraph()
    p_sign = doc.add_paragraph()
    p_sign.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_sign = p_sign.add_run(
        f"{identite['prenom']} {identite['nom']}"
    )
    _set_font(run_sign, 11, bold=True)

    Path(chemin_sortie).parent.mkdir(parents=True, exist_ok=True)
    doc.save(chemin_sortie)
    return chemin_sortie
